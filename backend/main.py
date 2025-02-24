from fastapi import FastAPI, UploadFile, File, HTTPException, Response, Header, Depends, APIRouter, Request
from fastapi import Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
import uvicorn
import os
import tempfile
import logging
from typing import Optional, Literal, List, Dict, Any
from langchain_ollama import OllamaLLM
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from google.cloud import translate_v2 as translate
from sentence_transformers import SentenceTransformer, util
from pydantic import BaseModel, ValidationError
from datetime import datetime
import ssl
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from transformers import MBartForConditionalGeneration, MBart50TokenizerFast
import io
import uuid
from fastapi.security import APIKeyHeader
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import pandas as pd
import matplotlib.pyplot as plt
import base64
from document_processing import extract_text, analyze_with_llm
import hashlib
from utils.text_extraction import extract_text
from utils.summary_generation import generate_summary
import json
from api.prompts import PromptCreate, app as prompts_router
from fastapi.exceptions import RequestValidationError
from database import init_db
import glob
from pathlib import Path
import shutil
import numpy as np
from azure_model_service import azure_service, get_azure_service
from langchain.chat_models import AzureChatOpenAI
from auth import get_current_user, User  # 显式导入User类
from fastapi.security import OAuth2PasswordBearer
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

logging.basicConfig(
    level=logging.DEBUG,  # 显示更详细日志
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('app.log')
    ]
)
logger = logging.getLogger(__name__)

app = FastAPI()
app.include_router(prompts_router)  # 挂载路由

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有来源
    allow_credentials=True,
    allow_methods=["*"],   # 允许所有方法
    allow_headers=["*"],   # 允许所有头
    expose_headers=["Content-Disposition"]  # 暴露必要头信息
)

# 获取项目根目录
BASE_DIR = Path(__file__).parent.resolve()
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)  # 确保目录存在

print(f"✅ 上传目录已创建：{UPLOAD_DIR.absolute()}")

# 初始化翻译模型
mbart_tokenizer = MBart50TokenizerFast.from_pretrained("facebook/mbart-large-50-many-to-many-mmt")
mbart_model = MBartForConditionalGeneration.from_pretrained("facebook/mbart-large-50-many-to-many-mmt")


LANGUAGE_CODES = {
    'zh': 'zh_CN',
    'en': 'en_XX'
}


mistral = OllamaLLM(
    base_url='http://localhost:11434',
    model="mistral:latest",
    temperature=0.7
)

llama = OllamaLLM(
    base_url='http://localhost:11434',
    model="llama3.2-vision:11b",
    temperature=0.7
)

def init_translation_model():
    try:
        logger.info("开始初始化翻译模型...")
        cache_dir = os.path.expanduser("~/.cache/huggingface/hub")
        model_name = "facebook/mbart-large-50-many-to-many-mmt"
        
        if torch.backends.mps.is_available():
            device = torch.device("mps")
            logger.info("使用 MPS 设备")
        elif torch.cuda.is_available():
            device = torch.device("cuda")
            logger.info("使用 CUDA 设备")
        else:
            device = torch.device("cpu")
            logger.info("使用 CPU 设备")
        
        logger.info("正在加载tokenizer...")
        tokenizer = AutoTokenizer.from_pretrained(
            model_name,
            use_fast=True,
            cache_dir=cache_dir
        )
        logger.info("tokenizer加载完成")
        
        logger.info("正在加载模型...")
        model = AutoModelForSeq2SeqLM.from_pretrained(
            model_name,
            torch_dtype=torch.float32,
            cache_dir=cache_dir,
            device_map="auto" if device.type != "cpu" else None,
            low_cpu_mem_usage=True
        )
        logger.info("模型加载完成")
        
        if device.type == "cpu":
            model = model.to(device)
        logger.info(f"模型已移动到 {device} 设备")
        
        return tokenizer, model, device
        
    except Exception as e:
        logger.error(f"模型初始化失败: {str(e)}")
        import traceback
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"模型加载错误: {str(e)}")

def process_pdf(file_path, tokenizer, translation_model, device):
    try:
        logger.info(f"开始处理PDF文件: {file_path}")
        
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        
        logger.info("正在加载PDF文件...")
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        logger.info(f"PDF加载完成，共 {len(pages)} 页")
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", "。", ".", "；", ";", "，", ",", "！", "!", "？", "?"]
        )
        
        all_chunks = []
        original_texts = []
        translated_texts = []
        
        for i, page in enumerate(pages):
            text = page.page_content.strip()
            if not text:
                continue
                
            page_chunks = text_splitter.split_text(text)
            
            for chunk in page_chunks:
                all_chunks.append({
                    'page': i + 1,
                    'content': chunk
                })
        
        logger.info(f"文本分割完成，共 {len(all_chunks)} 个片段")
        
        base_filename = os.path.basename(file_path)
        if base_filename.endswith('.pdf'):
            base_filename = base_filename[:-4]
            
        txt_file = os.path.join(UPLOAD_DIR, f"translated_{base_filename}.txt")
        docx_file = os.path.join(UPLOAD_DIR, f"translated_{base_filename}.docx")
        
        doc = Document()
        doc.add_heading('文档翻译结果', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(txt_file):
            os.remove(txt_file)
        
        successful_translations = 0
        for i, chunk in enumerate(all_chunks):
            try:
                if not chunk['content'].strip():
                    continue
                    
                logger.info(f"正在处理第 {i+1}/{len(all_chunks)} 个片段")
                logger.info(f"页码: {chunk['page']}")
                
                original_texts.append(chunk['content'])
                
                translated_text = translate_text(
                    chunk['content'],
                    tokenizer,
                    translation_model,
                    device
                )
                
                if not translated_text:
                    continue
                
                translated_texts.append(translated_text)
                
                with open(txt_file, "a", encoding="utf-8") as f:
                    f.write(f"\n{'='*50}\n")
                    f.write(f"页码: {chunk['page']}\n")
                    f.write(f"段落: {i+1}/{len(all_chunks)}\n")
                    f.write(f"原文：\n{chunk['content']}\n\n")
                    f.write(f"译文：\n{translated_text}\n")
                
                doc.add_heading(f'第{chunk["page"]}页 - 段落{i+1}', level=1)
                doc.add_heading('原文:', level=2)
                doc.add_paragraph(clean_text_for_docx(chunk['content']))
                doc.add_heading('译文:', level=2)
                doc.add_paragraph(clean_text_for_docx(translated_text))
                doc.add_paragraph('_' * 50)
                
                successful_translations += 1
                
            except Exception as e:
                logger.error(f"处理段落 {i+1} 时出错: {str(e)}")
                continue
        
        if successful_translations == 0:
            logger.error("没有成功翻译任何内容")
            return False, None, None, None
        
        try:
            doc.save(docx_file)
            logger.info("文档保存成功")
            
            if os.path.exists(txt_file) and os.path.exists(docx_file):
                txt_size = os.path.getsize(txt_file)
                docx_size = os.path.getsize(docx_file)
                
                if txt_size > 0 and docx_size > 0:
                    logger.info(f"文件生成成功: TXT={txt_size}字节, DOCX={docx_size}字节")
                    return True, base_filename, "\n".join(original_texts), "\n".join(translated_texts)
                    
            logger.error("文件验证失败")
            return False, None, None, None
            
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            return False, None, None, None
            
    except Exception as e:
        logger.error(f"PDF处理错误: {str(e)}")
        return False, None, None, None

def detect_language_and_direction(text):
    has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text)
    has_english = bool(re.search('[a-zA-Z]', text))
    
    if has_chinese and not has_english:
        return "zh2en"
    elif has_english and not has_chinese:
        return "en2zh"
    elif has_chinese and has_english:
        chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
        english_chars = sum(1 for char in text if char.isascii())
        return "zh2en" if chinese_chars > english_chars else "en2zh"
    else:
        return "en2zh"

def clean_text_for_docx(text):
    if not text:
        return ""
    return ''.join(char for char in text if char >= ' ' or char in ['\n', '\t'])

def translate_text(text: str, tokenizer, model, device="cpu"):
    try:
        logger.info("开始翻译文本...")
        logger.info(f"输入文本长度: {len(text)}")
        logger.info(f"使用设备: {device}")
        
        direction = detect_language_and_direction(text)
        logger.info(f"检测到的翻译方向: {direction}")
        
        if direction == "zh2en":
            src_lang = "zh_CN"
            tgt_lang = "en_XX"
        else:
            src_lang = "en_XX"
            tgt_lang = "zh_CN"
        
        logger.info(f"源语言: {src_lang}, 目标语言: {tgt_lang}")
        
        try:
            tokenizer.src_lang = src_lang
            inputs = tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)
            logger.info("文本标记化完成")
            
            inputs = {k: v.to(device) for k, v in inputs.items()}
            logger.info("输入已移动到设备")
            
            with torch.no_grad():
                logger.info("开始生成翻译...")
                outputs = model.generate(
                    **inputs,
                    forced_bos_token_id=tokenizer.lang_code_to_id[tgt_lang],
                    max_length=1024,
                    num_beams=5,
                    length_penalty=1.2,
                    no_repeat_ngram_size=3
                )
                logger.info("翻译生成完成")
            
            translated = tokenizer.decode(outputs[0], skip_special_tokens=True)
            logger.info(f"翻译完成，输出文本长度: {len(translated)}")
            
            return translated
            
        except Exception as e:
            logger.error(f"翻译过程中出错: {str(e)}")
            logger.error(f"错误类型: {type(e)}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
            raise
        
    except Exception as e:
        logger.error(f"翻译函数出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"翻译错误: {str(e)}")

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    tmp_path = None
    try:
        logger.info(f"开始处理文件: {file.filename}")
        
        if not file.filename.lower().endswith('.pdf'):
            logger.error("不支持的文件类型")
            raise HTTPException(status_code=400, detail="只支持PDF文件")
        
        try:
            content = await file.read()
            file_size = len(content)
            logger.info(f"文件大小: {file_size / (1024*1024):.2f} MB")
            
            if file_size == 0:
                logger.error("文件为空")
                raise HTTPException(status_code=400, detail="文件为空")
                
            if file_size > 500 * 1024 * 1024:
                logger.error("文件太大")
                raise HTTPException(status_code=400, detail="文件大小超过500MB")
                
        except Exception as e:
            logger.error(f"读取文件失败: {str(e)}")
            raise HTTPException(status_code=400, detail="文件读取失败")
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
                logger.info(f"临时文件已创建: {tmp_path}")
        except Exception as e:
            logger.error(f"创建临时文件失败: {str(e)}")
            raise HTTPException(status_code=500, detail="无法创建临时文件")
            
        try:
            logger.info("开始初始化翻译模型...")
            tokenizer, translation_model, device = init_translation_model()
            logger.info(f"翻译模型初始化完成，使用设备: {device}")
            
            logger.info("开始处理PDF文件...")
            success, base_filename, original_text, translated_text = process_pdf(
                tmp_path,
                tokenizer,
                translation_model,
                device
            )
            
            if not success:
                raise HTTPException(status_code=500, detail="PDF处理失败")
            
            docx_filename = f"translated_{base_filename}.docx"
            txt_filename = f"translated_{base_filename}.txt"
            
            docx_path = os.path.join(UPLOAD_DIR, docx_filename)
            txt_path = os.path.join(UPLOAD_DIR, txt_filename)
            
            if not os.path.exists(docx_path):
                logger.error(f"DOCX文件不存在: {docx_path}")
                raise HTTPException(status_code=500, detail="DOCX文件生成失败")
            
            if not os.path.exists(txt_path):
                logger.error(f"TXT文件不存在: {txt_path}")
                raise HTTPException(status_code=500, detail="TXT文件生成失败")
            
            logger.info("文件处理成功完成")
            return JSONResponse(content={
                "message": "文件处理成功",
                "files": {
                    "docx": docx_filename,
                    "txt": txt_filename
                },
                "original_text": original_text,
                "translated_text": translated_text
            })
            
        except Exception as e:
            logger.error(f"处理过程中出错: {str(e)}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=str(e))
            
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                    logger.info("临时文件已清理")
                except Exception as e:
                    logger.error(f"清理临时文件失败: {str(e)}")
                
    except Exception as e:
        logger.error(f"上传处理失败: {str(e)}")
        import traceback
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/downloads/{filename}")
async def download_file(filename: str):
    try:
        file_path = os.path.join(UPLOAD_DIR, filename)
        logger.info(f"尝试下载文件: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise HTTPException(status_code=404, detail="文件不存在")
        
        content_type = ('application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                       if filename.endswith('.docx') else 'text/plain')
        
        with open(file_path, 'rb') as f:
            file_content = f.read()
            
        return Response(
            content=file_content,
            media_type=content_type,
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Access-Control-Expose-Headers': 'Content-Disposition'
            }
        )
    except Exception as e:
        logger.error(f"下载文件时出错: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/debug")
async def debug_log(message: dict):
    log_type = message.get('type', 'info')
    log_msg = message.get('message', '')
    log_data = message.get('data', {})
    log_text = message.get('text', '')
    log_error = message.get('error', '')
    
    log_str = f"[前端日志] [{log_type.upper()}] {log_msg}"
    if log_text:
        log_str += f"\n文本: {log_text}"
    if log_data:
        log_str += f"\n数据: {log_data}"
    if log_error:
        log_str += f"\n错误: {log_error}"
    
    log_str = f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')}] {log_str}"
    
    if log_type == 'error':
        logger.error(log_str)
    else:
        logger.info(log_str)
    
    return {"status": "ok"}

class TranslationRequest(BaseModel):
    text: str
    mode: str = "fast"
    model: str = "local"
    llm_model: str = "mistral"
    direction: str  # 必须包含的字段

class TranslationResponse(BaseModel):
    translatedText: str
    confidence: float
    model_used: str

class TextGenerationRequest(BaseModel):
    text: str
    template: str = "default"  
    llm_model: str = None  # 添加模型选择

class TextGenerationResponse(BaseModel):
    generatedText: str
    model_used: str

@app.post("/api/translate")
async def translate(request: TranslationRequest):
    try:
        logger.info(f"收到翻译请求: {request}")
        
        # 根据模式调整参数
        if request.mode == "professional":
            # 使用更复杂的提示词
            prompt = f"请以专业医学翻译标准翻译以下内容：{request.text}"
        else:
            # 快速翻译的简单提示词
            prompt = f"翻译以下内容：{request.text}"
        
        src_lang = 'zh_CN' if request.direction == 'zh2en' else 'en_XX'
        tgt_lang = 'en_XX' if request.direction == 'zh2en' else 'zh_CN'
        
        MAX_LENGTH = 400  
        text_segments = []
        current_text = prompt
        
        current_text = current_text.replace('\r\n', '\n').replace('\r', '\n')
        
        sections = current_text.split('\r')
        current_segment = ""
        
        for section in sections:
            items = section.split('\n')
            for item in items:
                if not item.strip():
                    continue
                
                if item.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
                    if current_segment:
                        text_segments.append(current_segment.strip())
                    current_segment = item
                    continue
                
                if len(current_segment) + len(item) > MAX_LENGTH:
                    if current_segment:
                        text_segments.append(current_segment.strip())
                    current_segment = item
                else:
                    if current_segment:
                        current_segment += '\n' + item
                    else:
                        current_segment = item
        
        if current_segment:
            text_segments.append(current_segment.strip())
        
        logger.info(f"分段数量: {len(text_segments)}")
        for i, seg in enumerate(text_segments):
            logger.info(f"段落 {i+1} 长度: {len(seg)}")
        
        translated_segments = []
        
        for segment in text_segments:
            mbart_tokenizer.src_lang = src_lang
            
            encoded = mbart_tokenizer(segment, return_tensors="pt", padding=True)
            
            generated_tokens = mbart_model.generate(
                **encoded,
                forced_bos_token_id=mbart_tokenizer.lang_code_to_id[tgt_lang],
                max_length=1024,
                num_beams=5,
                length_penalty=1.0,
                early_stopping=True
            )
            
            segment_translation = mbart_tokenizer.batch_decode(generated_tokens, skip_special_tokens=True)[0]
            translated_segments.append(segment_translation)
        
        translated_text = '\n\n'.join(translated_segments) 
        
        logger.info(f"翻译结果: {translated_text}")
        return TranslationResponse(
            translatedText=translated_text,
            confidence=0.95,
            model_used="mbart-large-50"
        )
    except Exception as e:
        logger.error(f"翻译错误: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

class QuestionRequest(BaseModel):
    question: str
    file_path: str
    temperature: float = 0.7

@app.post("/api/qa")
async def question_answer(request: QuestionRequest):
    try:
        logger.info(f"收到问答请求: {request}")
        
        file_path = os.path.join(UPLOAD_DIR, request.file_path)
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文档不存在")
            
        with open(file_path, 'r', encoding='utf-8') as f:
            document_text = f.read()
            
        ollama = OllamaLLM(base_url='http://localhost:11434', model="mistral:latest")
        
        prompt = f"""你是一个专业的医学文献问答助手。请基于以下医学文献内容，回答用户的问题。
如果文献中没有相关信息，请直接说明。

医学文献内容：
{document_text}

用户问题：{request.question}

请用专业、准确且易懂的语言回答问题。如果文献中没有相关信息，请回答"抱歉，文献中没有找到相关信息。"
"""
        
        try:
            response = ollama.invoke(prompt, temperature=request.temperature)
            logger.info(f"生成的回答: {response}")
            
            return {"answer": response}
            
        except Exception as e:
            logger.error(f"生成回答时出错: {str(e)}")
            raise HTTPException(status_code=500, detail=f"生成回答失败: {str(e)}")
        
    except Exception as e:
        logger.error(f"问答失败: {str(e)}")
        import traceback
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

def init_google_translate():
    try:
        os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = CREDENTIALS_PATH
        translate_client = translate.Client()
        return translate_client
    except Exception as e:
        logger.error(f"Google翻译客户端初始化失败: {str(e)}")
        raise

def back_translate(text: str, translate_client):
    try:
        detection = translate_client.detect_language(text)
        source_lang = detection['language']
        logger.info(f"检测到的原文语言: {source_lang}")
        
        intermediate_lang = 'en' if source_lang == 'zh' else 'zh'
        target_lang = source_lang
        
        logger.info(f"翻译方向: {source_lang} -> {intermediate_lang} -> {target_lang}")
        
        intermediate_result = translate_client.translate(
            text,
            target_language=intermediate_lang,
            source_language=source_lang
        )
        
        final_result = translate_client.translate(
            intermediate_result['translatedText'],
            target_language=target_lang,
            source_language=intermediate_lang
        )
        
        logger.info("反向翻译完成")
        return {
            'original': text,
            'intermediate': intermediate_result['translatedText'],
            'final': final_result['translatedText']
        }
        
    except Exception as e:
        logger.error(f"反向翻译失败: {str(e)}")
        raise

class TextRequest(BaseModel):
    text: str

@app.post("/api/back-translate")
async def back_translate_api(request: TextRequest):
    try:
        translate_client = init_google_translate()
        result = back_translate(request.text, translate_client)
        return result
    except Exception as e:
        logger.error(f"反向翻译API调用失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def health_check():
    return {
        "status": "running",
        "upload_dir": str(UPLOAD_DIR.absolute()),
        "files": [f.name for f in UPLOAD_DIR.glob("*")]
    }

@app.get("/test-connection")
async def test_connection():
    logger.info("收到测试连接请求")
    return {"status": "ok", "message": "连接成功"}

@app.post("/api/generate")
async def generate_text(request: TextGenerationRequest):
    try:
        logger.info(f"收到生成请求: {request}")
        
        model = llama if request.llm_model == 'llama' else mistral
        
        # 根据不同的模板选择不同的处理逻辑
        if request.template == 'compliance':
            # 合规检查的模板
            template = COMPLIANCE_TEMPLATE
        elif request.template == 'csr':
            # CSR生成的模板
            template = CSR_TEMPLATE
        else:
            raise HTTPException(status_code=400, detail="Unsupported template type")
        
        prompt = template.format(text=request.text)
        
        try:
            # 尝试使用MPS加速
            model = load_local_model('mistral', device='mps')
        except RuntimeError as e:
            if 'MPS not available' in str(e):
                # 回退到CPU
                model = load_local_model('mistral', device='cpu')
                logger.warning("MPS不可用，已回退到CPU模式")
            else:
                raise
        
        response = model(prompt)
        
        return TextGenerationResponse(
            generatedText=response.strip(),
            model_used=request.llm_model or "mistral"
        )
        
    except Exception as e:
        logger.error(f"生成失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# 定义合规检查模板
COMPLIANCE_TEMPLATE = {
    "id": "fda-compliance",
    "title": "FDA Compliance Check",
    "content": """Please analyze the following text for regulatory compliance:

Text to analyze: {text}

Please provide your analysis in the following structured format:

1. Understanding the Context and Scope:
• Provide a brief overview of the document's context
• Identify the type of submission and regulatory framework

2. Identifying Gaps in Provided Documents:
• List each identified gap using lettered sub-points (a, b, c, etc.)
• For each gap, cite the specific regulatory requirement it fails to meet

3. Addressing Specific Areas of Focus:
• Pharmacokinetics: Required data and current status
• Toxicology: Required studies and current status
• Safety Pharmacology: Required studies and current status

4. FDA-Specific Considerations:
• List specific FDA requirements
• Identify any FDA-specific gaps""",
    "isDefault": True,
    "modelType": "compliance",
    "category": "regulatory"
}

# 定义 CSR 生成模板
CSR_TEMPLATE = {
    "id": "csr-generation",
    "title": "CSR Generation (NDA)",
    "content": """Generate a Clinical Study Report (CSR) for an FDA NDA submission based on the content extracted from the provided file.

Content to analyze: {text}

Please generate a comprehensive Clinical Study Report (CSR) with the following structure:

1. Structure & Table of Contents
• Title Page
• Synopsis
• Table of Contents (detailing every major section)
• List of Abbreviations and Definitions of Terms
• Ethics
• IRB/IEC approvals
• Ethical conduct statement
• Investigators and Study Administrative Structure
• Introduction
• Study Objectives
• Primary Objectives
• Secondary Objectives
• Exploratory Objectives (if applicable)
• Investigational Plan
• Study Design
• Study Population
• Treatment/Study Medication Details
• Efficacy Evaluation
• Safety Evaluation
• Statistical Methods
• Sample Size Rationale
• Analysis Methods for Efficacy and Safety
• Results
• Participant Disposition
• Demographics and Baseline Characteristics
• Efficacy Results
• Safety Results (including AEs/SAEs)
• Discussion and Conclusions
• References
• Appendices (e.g., Protocol, Sample CRFs, Patient Data Listings, etc.)

Note: Please maintain professional, scientific language and present data objectively. Include both positive and negative findings where available.""",
    "isDefault": True,
    "modelType": "generation",
    "category": "csr"
}

@app.post("/api/download-report")
async def download_report(request: dict):
    try:
        content = request.get('content')
        file_name = request.get('fileName')
        
        doc = Document()
        
        title = doc.add_heading('Compliance Check Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(datetime.now().strftime("%Y-%m-%d")).italic = True
        
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('  '):  
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.add_run(line.strip())
                else:  
                    doc.add_paragraph(line)
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return Response(
            content=f.read(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{file_name}"'
            }
        )
        
    except Exception as e:
        logger.error(f"生成报告失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/log")
async def log_message(request: dict):
    try:
        logger.info(f"Add-in log: {request.get('message')}")
        if request.get('data'):
            logger.info(f"Data: {request.get('data')}")
        return {"status": "success"}
    except Exception as e:
        logger.error(f"Failed to log message: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# 在文件顶部添加预定义提示词
PREDEFINED_PROMPTS = [CSR_TEMPLATE, COMPLIANCE_TEMPLATE]

# 修改获取提示词的API端点
@app.get("/api/prompts")
async def get_prompts():
    return JSONResponse({
        "defaultPrompts": PREDEFINED_PROMPTS,  
        "userPrompts": get_user_prompts_from_db()
    })

@app.middleware("http")
async def validate_api_key(request: Request, call_next):
    # 允许预检请求
    if request.method == "OPTIONS":
        return await call_next(request)
        
    # 排除文档路由
    if request.url.path in ['/docs', '/openapi.json', '/redoc']:
        return await call_next(request)
    
    api_key = request.headers.get('X-API-Key')
    valid_key = "your-secure-key-here"  # 替换为实际密钥
    
    if not api_key or api_key != valid_key:
        return JSONResponse(
            status_code=403,
            content={"detail": "Invalid or missing API key"},
            headers={"Access-Control-Allow-Origin": "*"}
        )
    return await call_next(request)

class ExecutionRequest(BaseModel):
    text: str
    prompt_id: str
    model: Literal['mistral', 'llama']

@app.post("/api/execute")
async def execute_prompt(request: ExecutionRequest):
    # 添加详细验证
    if not request.text.strip():
        logger.error("执行请求缺少文本内容")
        raise HTTPException(status_code=422, detail="Text content is required")
    
    if not request.prompt_id.strip():
        logger.error("执行请求缺少提示词ID")
        raise HTTPException(status_code=422, detail="Prompt ID is required")
    
    # 获取完整提示词
    prompt = next((p for p in prompt_storage["default"] + prompt_storage["users"].get("demo-user", []) 
                  if p["id"] == request.prompt_id), None)
    
    if not prompt:
        raise HTTPException(status_code=404, detail="Prompt not found")
    
    # 构建完整提示
    full_prompt = prompt["content"].replace("{text}", request.text)
    
    # 根据模型类型调用不同处理
    if request.model == 'llama':
        result = compliance_check(full_prompt)
    else:
        result = generate_content(full_prompt)
    
    return {"result": result}

# 添加生成函数
def generate_content(prompt: str) -> str:
    """调用LLM生成内容"""
    try:
        llm = OllamaLLM(model="mistral")
        return llm.invoke(prompt)
    except Exception as e:
        logging.error(f"生成内容失败: {str(e)}")
        raise HTTPException(status_code=500, detail="生成失败")

# 添加合规检查函数
def compliance_check(prompt: str) -> str:
    """调用LLM进行合规检查"""
    try:
        llm = OllamaLLM(model="llama3.2-vision:11b")
        return llm.invoke(prompt)
    except Exception as e:
        logging.error(f"合规检查失败: {str(e)}")
        raise HTTPException(status_code=500, detail="合规检查失败")

# 新增属性模型
class DocumentAttribute(BaseModel):
    name: str
    value: str
    editable: bool

class TemplateAttribute(DocumentAttribute):
    applicable_models: List[str]

# 新增API端点
@app.post("/api/save-attributes")
async def save_attributes(
    doc_attrs: List[DocumentAttribute],
    template_attrs: List[TemplateAttribute]
):
    try:
        # 实现属性存储逻辑
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# 新增数据模型
class Source(BaseModel):
    id: str
    name: str
    type: str = "document"  # 默认类型
    origin: str = "upload"  # 默认来源
    size: str
    isTemplate: bool = False
    summary: str = ""

class Template(BaseModel):
    id: str
    name: str
    category: Literal['clinical', 'regulatory', 'csr']

# 新增API端点
source_db = [
    {"id": "doc-001", "name": "Main Protocol", "type": "document", "origin": "internal"},
    {"id": "cro-2023", "name": "CRO Final Report 2023", "type": "report", "origin": "external"},
    {"id": "lab-456", "name": "Lab Results Q4", "type": "lab", "origin": "external"}
]

template_db = [
    {"id": "clin-001", "name": "Clinical Study Report", "category": "clinical"},
    {"id": "reg-2023", "name": "EU MDR Template", "category": "regulatory"},
    {"id": "csr-01", "name": "CSR Brief", "category": "csr"}
]

@app.get("/api/sources")
async def get_all_sources():
    try:
        document_files = glob.glob(os.path.join(UPLOAD_DIR, "*"))
        return [
            {
                "id": os.path.splitext(os.path.basename(f))[0],  # 使用文件名作为ID
                "name": os.path.basename(f),  # 返回实际文件名
                "type": "document",
                "origin": "upload",
                "size": str(os.path.getsize(f)),
                "uploadDate": datetime.fromtimestamp(
                    os.path.getctime(f)
                ).isoformat(),
                "vectorized": True,
                "analyzed": False
            }
            for f in document_files
            if not os.path.basename(f).startswith('.')  # 排除隐藏文件
        ]
    except Exception as e:
        logger.error(f"获取文档列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/templates", response_model=List[Template])
async def get_templates():
    return template_db

# 添加数据模型
class VisualizationConfig(BaseModel):
    chartType: str
    xAxis: str
    yAxis: str

class DatasetInfo(BaseModel):
    data: List[Dict[str, Any]]
    columns: List[str]

class VisualizationRequest(BaseModel):
    dataset: DatasetInfo
    config: VisualizationConfig

# 添加API端点
dataset_storage = {}

@app.post("/api/datasets/upload")
async def upload_dataset(file: UploadFile):
    try:
        logger.info(f"Received dataset upload: {file.filename}")
        content = await file.read()
        logger.info(f"File content read: {len(content)} bytes")
        
        # 使用 pandas 读取 CSV，添加数据类型处理
        df = pd.read_csv(
            io.StringIO(content.decode('utf-8')),
            low_memory=False,  # 防止混合类型警告
            na_values=['NA', 'N/A', ''],  # 处理缺失值
            dtype_backend='numpy_nullable'  # 使用新的 dtype 后端处理混合类型
        )
        
        # 处理无限和非法浮点数值
        for column in df.select_dtypes(include=['float64']).columns:
            # 将无限值替换为 NaN
            df[column] = df[column].replace([np.inf, -np.inf], np.nan)
            # 将 NaN 替换为 None (会被转换为 JSON 中的 null)
            df[column] = df[column].where(pd.notnull(df[column]), None)

        logger.info(f"Dataset loaded: {len(df)} rows, {len(df.columns)} columns")
        
        # 确保返回正确的列信息和数据格式
        return {
            "message": "Dataset uploaded successfully",
            "rows": len(df),
            "columns": df.columns.tolist(),  # 返回列名列表
            "data": df.head(100).to_dict('records'),  # 返回前100行数据
            "preview": df.head(5).to_dict('records')  # 返回前5行预览
        }
        
    except Exception as e:
        logger.error(f"Dataset upload failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process dataset: {str(e)}"
        )

@app.post("/api/generate-visualization")
async def generate_visualization(config: dict):
    try:
        logger.info(f"Generating visualization with config: {config}")
        
        # 获取当前数据集
        df = pd.DataFrame(current_dataset)
        
        # 验证数据列
        x_col = config.get('xAxis')
        y_col = config.get('yAxis')
        chart_type = config.get('chartType', 'bar')
        
        if not x_col or not y_col:
            raise HTTPException(status_code=400, detail="Missing axis configuration")
            
        # 自动转换数值类型
        if not np.issubdtype(df[y_col].dtype, np.number):
            try:
                df[y_col] = pd.to_numeric(df[y_col], errors='coerce')
                df = df.dropna(subset=[y_col])
                if df.empty:
                    raise ValueError("No valid numeric data after conversion")
            except Exception as e:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Column {y_col} cannot be converted to numeric values"
                )
        
        # 生成图表
        plt.figure(figsize=(10, 6))
        if chart_type == 'bar':
            df.plot.bar(x=x_col, y=y_col, ax=plt.gca())
        elif chart_type == 'line':
            df.plot.line(x=x_col, y=y_col, ax=plt.gca())
        elif chart_type == 'pie':
            df[y_col].value_counts().plot.pie(autopct='%1.1f%%', ax=plt.gca())
        else:
            raise HTTPException(status_code=400, detail="Unsupported chart type")
            
        # 保存图表
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', bbox_inches='tight')
        plt.close()
        img_buffer.seek(0)
        
        return StreamingResponse(img_buffer, media_type="image/png")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to generate visualization: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

chat_router = APIRouter()

class ChatMessage(BaseModel):
    role: str  # 'user' or 'assistant'
    content: str

class ChatRequest(BaseModel):
    question: str
    document_text: str
    history: List[ChatMessage]
    model: str = "local"

@app.post("/api/chat/word")
async def chat_with_word(request: ChatRequest):
    try:
        logger.info(f"[Chat] Processing request with history length: {len(request.history)}")
        
        # 模型选择逻辑
        use_cloud = request.model == "azure"
        logger.info(f"Using {'Azure' if use_cloud else 'local'} model: {request.model}")
        
        if not request.history:
            logger.info("[Chat] No history provided, starting new conversation")
        else:
            logger.info(f"[Chat] Last message in history: {request.history[-1].content[:100]}...")

        # 格式化聊天历史
        formatted_history = [
            (msg.content, response.content)
            for msg, response in zip(
                [m for m in request.history if m.role == "user"],
                [m for m in request.history if m.role == "assistant"]
            )
        ]

        # 处理文档
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_text(request.document_text)
        
        # 创建临时向量库
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-mpnet-base-v2",
            model_kwargs={'device': 'mps'}  # 使用M1/M2 Mac的MPS加速
        )
        vectorstore = FAISS.from_texts(chunks, embeddings)

        # 模型选择核心逻辑
        if use_cloud:
            # 使用Azure云服务
            service = await get_azure_service()
            qa = ConversationalRetrievalChain.from_llm(
                llm=AzureChatOpenAI(
                    azure_deployment=os.getenv("AZURE_ENGINE"),
                    openai_api_version=os.getenv("AZURE_API_VERSION")
                ),
                retriever=vectorstore.as_retriever(),
                return_source_documents=True
            )
        else:
            # 本地模型选择
            local_model = llama if request.model == "llama" else mistral
            qa = ConversationalRetrievalChain.from_llm(
                llm=local_model,
                retriever=vectorstore.as_retriever(),
                return_source_documents=True
            )

        # 执行对话（保持原有逻辑）
        result = await qa.ainvoke({
            "question": request.question,
            "chat_history": formatted_history,
            "context": request.document_text
        })

        return {"response": result["answer"]}
        
    except ValidationError as e:
        logger.error(f"[Chat] Validation error: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"[Chat] Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# 将路由挂载到主应用
app.include_router(chat_router, prefix="/api")

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class DocumentMetadata(BaseModel):
    id: str
    filename: str
    content_hash: str
    upload_time: datetime
    file_size: int
    summary: Optional[str] = None

# 文件信息持久化存储路径
FILES_DB_PATH = os.path.join(os.path.dirname(__file__), "files_db.json")

def load_files_db():
    if os.path.exists(FILES_DB_PATH):
        with open(FILES_DB_PATH, 'r') as f:
            logger.info("Loading files database")
            return json.load(f)
    logger.info("Creating new files database")
    return {}

def save_files_db(files_db):
    logger.info(f"Saving files database. Current files: {json.dumps(files_db, indent=2)}")
    with open(FILES_DB_PATH, 'w') as f:
        json.dump(files_db, f)

# 初始化文件数据库
stored_files = load_files_db()

def process_document(file_path: str):
    # 根据文件类型选择加载器
    if file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith('.docx'):
        loader = Docx2txtLoader(file_path)
    else:
        raise ValueError("Unsupported file format")

    documents = loader.load()
    
    # 文本分块
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_documents(documents)
    
    # 生成向量存储
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2",
        model_kwargs={'device': 'mps'}
    )
    vectorstore = FAISS.from_documents(chunks, embeddings)
    
    # 保存时确保设置正确的序列化选项
    vectorstore.save_local(
        "vectorstore",
        # 如果需要，可以添加其他序列化选项
    )
    
    return chunks

@app.post("/api/sources/upload")
async def upload_document(file: UploadFile = File(...)):
    try:
        # 生成唯一ID但保留原始文件名
        file_id = str(uuid.uuid4())
        original_name = file.filename
        file_ext = os.path.splitext(original_name)[1]
        
        # 保存文件，使用原始文件名
        file_path = os.path.join(UPLOAD_DIR, original_name)
        
        # 如果文件名已存在，添加数字后缀
        base_name = os.path.splitext(original_name)[0]
        counter = 1
        while os.path.exists(file_path):
            new_name = f"{base_name}_{counter}{file_ext}"
            file_path = os.path.join(UPLOAD_DIR, new_name)
            counter += 1
            
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        return {
            "document": {
                "id": file_id,
                "name": os.path.basename(file_path),
                "type": "document",
                "origin": "upload",
                "size": str(os.path.getsize(file_path)),
                "uploadDate": datetime.now().isoformat(),
                "vectorized": False,  # 初始上传时为false
                "analyzed": False     # 初始上传时为false
            }
        }
    except Exception as e:
        logger.error(f"Upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/sources/{doc_id}")
async def delete_source(doc_id: str):
    try:
        # 检查文件是否存在
        doc_path = os.path.join(UPLOAD_DIR, f"{doc_id}.*")
        files = glob.glob(doc_path)
        
        if not files:
            raise HTTPException(
                status_code=404,
                detail="Document not found in uploads directory"
            )
            
        # 删除物理文件
        for file_path in files:
            try:
                os.remove(file_path)
                logger.info(f"Deleted file: {file_path}")
            except Exception as e:
                logger.error(f"Failed to delete file {file_path}: {str(e)}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to delete file: {str(e)}"
                )
        
        return {"status": "success"}
        
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Delete error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def is_template_document(doc_id: str) -> bool:
    # Add logic to check if document is a template
    template_ids = ["template1", "template2", "template3"]  # Store these in config
    return doc_id in template_ids

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request, exc):
    return JSONResponse(
        status_code=422,
        content={
            "detail": "Validation Error",
            "errors": exc.errors()
        },
    )

# 在FastAPI app创建后初始化数据库
init_db()

@app.post("/api/analyze")
async def analyze_documents(request: Request):
    data = await request.json()
    doc_ids = data.get("docIds", [])
    
    try:
        response = {
            "documents": []
        }
        
        for doc_id in doc_ids:
            doc_path = os.path.join(UPLOAD_DIR, f"{doc_id}.*")
            files = glob.glob(doc_path)
            if files:
                file_path = files[0]
                # 在这里进行向量化处理
                # ...
                
                response["documents"].append({
                    "id": doc_id,
                    "name": os.path.basename(file_path),
                    "type": "document",
                    "origin": "upload",
                    "size": str(os.path.getsize(file_path)),
                    "uploadDate": datetime.fromtimestamp(
                        os.path.getctime(file_path)
                    ).isoformat(),
                    "vectorized": True,   # 分析后设置为true
                    "analyzed": True      # 分析后设置为true
                })
        
        return response
        
    except Exception as e:
        logger.error(f"分析文档时出错: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def init_uploaded_files():
    """启动时加载已上传文件"""
    UPLOAD_DIR = "uploads"
    if not os.path.exists(UPLOAD_DIR):
        return []
    
    return [
        {
            "id": os.path.splitext(f)[0],  # 使用文件名作为ID
            "name": f,
            "type": "document",
            "origin": "upload",
            "uploadDate": datetime.fromtimestamp(
                os.path.getctime(os.path.join(UPLOAD_DIR, f))
            ).isoformat()
        }
        for f in os.listdir(UPLOAD_DIR)
    ]

@app.on_event("startup")
async def startup_event():
    if os.getenv("USE_CLOUD_MODELS", "false").lower() == "true":
        await azure_service.initialize()
    init_db()
    # 确保上传目录存在
    UPLOAD_DIR = "uploads"
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    logger.info("✅ 服务启动完成")
    logger.info(f"当前工作目录：{os.getcwd()}")
    logger.info(f"上传目录内容：{os.listdir(UPLOAD_DIR)}")

@app.middleware("http")
async def log_requests(request: Request, call_next):
    logger.info(f"收到请求: {request.method} {request.url}")
    try:
        response = await call_next(request)
        logger.info(f"返回响应: {response.status_code}")
        return response
    except Exception as e:
        logger.error(f"请求处理失败: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"message": "Internal server error"}
        )

def sanitize_xml(text: str) -> str:
    # 移除控制字符（ASCII 0-31，除了换行和制表符）
    cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # 替换非法XML字符
    return cleaned.encode('utf-8', 'ignore').decode('utf-8')

@app.post("/api/export-chat")
async def export_chat(request: Request):
    try:
        logger.info("开始处理导出请求")
        data = await request.json()
        # 清洗文档内容
        doc_text = sanitize_xml(data.get('documentText', 'No document content'))
        
        # 创建文档
        doc = Document()
        doc.add_heading('Chat Export', 0)
        
        # 添加聊天记录（清洗每条消息）
        for msg in data['messages']:
            role = "User" if msg.get('isUser') else "AI"
            content = sanitize_xml(msg.get('content', 'No content'))
            doc.add_paragraph(f"{role}: {content}")
        
        # 添加文档内容
        if doc_text.strip():
            doc.add_heading('Current Document', 1)
            doc.add_paragraph(doc_text)
        
        # 保存文件
        filename = f"ChatExport_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        export_path = os.path.join(EXPORT_DIR, filename)
        doc.save(export_path)
        
        logger.info(f"文件已保存到：{export_path}")
        return FileResponse(export_path, filename=filename)
    except Exception as e:
        logger.error(f"导出失败详情：{traceback.format_exc()}")
        raise HTTPException(500, "导出失败，请检查日志")

EXPORT_DIR = "exports"
if not os.path.exists(EXPORT_DIR):
    os.makedirs(EXPORT_DIR)

@app.post("/api/switch-model")
async def switch_model(
    use_cloud: bool = Body(..., embed=True),
    current_user: User = Depends(get_current_user)
):
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Requires admin privileges")
    
    try:
        if use_cloud:
            await azure_service.initialize()
        # 可以在这里添加本地模型的关闭逻辑
        return {"status": "success", "using_cloud": use_cloud}
    except Exception as e:
        logger.error(f"Model switch failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to switch model mode: {str(e)}"
        )

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",  # 改为0.0.0.0而不是localhost
        port=8000,
        reload=True,
        # 暂时禁用SSL
        # ssl_keyfile="./key.pem",
        # ssl_certfile="./cert.pem",
        log_level="debug"  # 添加详细日志
    ) 